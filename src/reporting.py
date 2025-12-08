# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import sqlite3
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

from docx import Document  # pip install python-docx

from src.db import DB_PATH, init_db, RESULTS_DIR


# ---------------------------------------------------------------------------
# Модели и утилиты
# ---------------------------------------------------------------------------

@dataclass
class VersionConstraint:
    op: str  # "<", "<=", "==", ">=", ">", "range"
    v1: Tuple[int, ...]
    v2: Optional[Tuple[int, ...]] = None  # для диапазона


def normalize_text(s: str) -> List[str]:
    """
    Разбиваем строку на токены: буквы+цифры, в нижнем регистре.
    Никаких доменных стоп-слов — только базовая нормализация.
    """
    tokens = re.split(r'[^0-9A-Za-zА-Яа-я]+', s.lower())
    return [t for t in tokens if t]


def tokens_similarity(a: Sequence[str], b: Sequence[str]) -> float:
    """
    Простая метрика похожести: |пересечение| / max(|A|, |B|).
    """
    set_a, set_b = set(a), set(b)
    if not set_a or not set_b:
        return 0.0
    inter = len(set_a & set_b)
    return inter / max(len(set_a), len(set_b))


def parse_version(s: str) -> Optional[Tuple[int, ...]]:
    """
    Преобразует '7.0.14' -> (7, 0, 14).
    Если ничего похожего на версию нет — возвращает None.
    """
    if not s:
        return None
    parts = re.findall(r'\d+', s)
    if not parts:
        return None
    return tuple(int(p) for p in parts)


def compare_versions(a: Tuple[int, ...], b: Tuple[int, ...]) -> int:
    """
    Лексикографическое сравнение версий.
    Возвращает:
      -1 если a < b
       0 если a == b
      +1 если a > b
    """
    max_len = max(len(a), len(b))
    aa = a + (0,) * (max_len - len(a))
    bb = b + (0,) * (max_len - len(b))
    if aa < bb:
        return -1
    if aa > bb:
        return 1
    return 0


# ---------------------------------------------------------------------------
# Извлечение версионных ограничений из строки product
# ---------------------------------------------------------------------------

def extract_version_constraints(product: str) -> List[VersionConstraint]:
    """
    Пытаемся понять из текста product, какие версии уязвимы.
    Поддерживаем базовые варианты:
      - 'before 7.3.1'         -> op "<"
      - 'prior to 7.3.1'       -> op "<"
      - 'through 7.3.1'        -> op "<="
      - '<= 7.3.1' / '>= 7.3.1'
      - '7.3.1 and earlier'    -> op "<="
      - 'version 7.3.1'        -> op "==" (если ничего больше нет)
    Если ничего уверенного не нашли — возвращаем пустой список.
    """
    if not product:
        return []

    text = product.lower()
    constraints: List[VersionConstraint] = []

    # before / prior to
    for m in re.finditer(r'(before|prior to)\s+([0-9][0-9A-Za-z\.\-\_]*)', text):
        ver = parse_version(m.group(2))
        if ver:
            constraints.append(VersionConstraint('<', ver))

    # through
    for m in re.finditer(r'(through)\s+([0-9][0-9A-Za-z\.\-\_]*)', text):
        ver = parse_version(m.group(2))
        if ver:
            constraints.append(VersionConstraint('<=', ver))

    # <= / >=
    for m in re.finditer(r'(<=|>=|<|>)\s*([0-9][0-9A-Za-z\.\-\_]*)', text):
        op = m.group(1)
        ver = parse_version(m.group(2))
        if ver:
            constraints.append(VersionConstraint(op, ver))

    # 'X and earlier'
    for m in re.finditer(r'([0-9][0-9A-Za-z\.\-\_]*)\s+and\s+earlier', text):
        ver = parse_version(m.group(1))
        if ver:
            constraints.append(VersionConstraint('<=', ver))

    # 'version 7.3.1' / 'versions 7.3.1'
    # Добавляем только если пока ничего другого не нашли.
    if not constraints:
        m = re.search(r'versions?\s+([0-9][0-9A-Za-z\.\-\_]*)', text)
        if m:
            ver = parse_version(m.group(1))
            if ver:
                constraints.append(VersionConstraint('==', ver))

    return constraints


def version_satisfies(installed: Tuple[int, ...], cons: VersionConstraint) -> bool:
    cmp_res = compare_versions(installed, cons.v1)
    if cons.op == '<':
        return cmp_res == -1
    if cons.op == '<=':
        return cmp_res in (-1, 0)
    if cons.op == '==':
        return cmp_res == 0
    if cons.op == '>':
        return cmp_res == 1
    if cons.op == '>=':
        return cmp_res in (0, 1)
    if cons.op == 'range' and cons.v2 is not None:
        return compare_versions(installed, cons.v1) >= 0 and compare_versions(installed, cons.v2) <= 0
    return False


def version_is_vulnerable(installed_version: Optional[str], product_text: str) -> str:
    """
    Возвращает один из статусов:
      - 'match'     — версия попадает в диапазон уязвимых
      - 'mismatch'  — явное несоответствие (например, установленная > 'before X')
      - 'unknown'   — не удалось однозначно оценить (нет версии или нет явных ограничений)
    """
    if not installed_version:
        return 'unknown'

    iv = parse_version(installed_version)
    if not iv:
        return 'unknown'

    constraints = extract_version_constraints(product_text)
    if not constraints:
        # нет явных ограничений — не делаем вывод, считаем неизвестным
        return 'unknown'

    any_match = any(version_satisfies(iv, c) for c in constraints)
    any_contradiction = any(not version_satisfies(iv, c) for c in constraints)

    if any_match:
        return 'match'
    if any_contradiction:
        return 'mismatch'
    return 'unknown'


# ---------------------------------------------------------------------------
# Загрузка данных из БД
# ---------------------------------------------------------------------------

def _load_data() -> Tuple[List[sqlite3.Row], Dict[str, List[sqlite3.Row]], List[sqlite3.Row]]:
    init_db()
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    try:
        agents = conn.execute('SELECT * FROM agents').fetchall()
        software_rows = conn.execute('SELECT * FROM agent_software').fetchall()
        vulns = conn.execute(
            '''
            SELECT id, cve, vendor, product, title, cvss, severity, url
            FROM vulnerabilities
            '''
        ).fetchall()
    finally:
        conn.close()

    soft_by_agent: Dict[str, List[sqlite3.Row]] = defaultdict(list)
    for row in software_rows:
        soft_by_agent[row['agent_id']].append(row)

    return agents, soft_by_agent, vulns


# ---------------------------------------------------------------------------
# Сопоставление ПО ↔ уязвимостей
# ---------------------------------------------------------------------------

SIM_THRESHOLD_NAME = 0.5   # порог похожести имени продукта
SIM_THRESHOLD_VENDOR = 0.5  # порог похожести вендора (если оба указаны)


def match_vulns_for_agent(
    agent: sqlite3.Row,
    soft_by_agent: Dict[str, List[sqlite3.Row]],
    vulns: List[sqlite3.Row],
) -> Dict[int, Dict[str, Any]]:
    """
    Возвращает словарь:
      vuln_id -> {
         'vuln': row_vuln,
         'software': [rows_software...]
      }
    Сопоставление:
      1) похожесть названия ПО и product (по токенам, % совпадения);
      2) если есть vendor/publisher — доп.критерий;
      3) попытка проверить версию (match / mismatch / unknown).
    """
    matches: Dict[int, Dict[str, Any]] = {}
    agent_soft = soft_by_agent.get(agent['agent_id'], [])

    for sw in agent_soft:
        sw_name = (sw['name'] or '').strip()
        if not sw_name:
            continue

        sw_publisher = (sw['publisher'] or '').strip()
        sw_name_tokens = normalize_text(sw_name)
        sw_pub_tokens = normalize_text(sw_publisher)

        for v in vulns:
            vid = int(v['id'])
            v_product = (v['product'] or '').strip()
            v_vendor = (v['vendor'] or '').strip()

            # 1. Похожесть названия продукта
            prod_tokens = normalize_text(v_product) or normalize_text(v['title'] or '')
            if not prod_tokens:
                continue

            sim_name = tokens_similarity(sw_name_tokens, prod_tokens)
            if sim_name < SIM_THRESHOLD_NAME:
                continue

            # 2. Похожесть вендора, если оба указаны
            if v_vendor and sw_publisher:
                vend_tokens = normalize_text(v_vendor)
                vend_sim = tokens_similarity(sw_pub_tokens, vend_tokens)
                if vend_sim < SIM_THRESHOLD_VENDOR:
                    continue

            # 3. Версионность
            vcheck = version_is_vulnerable(sw['version'], v_product)

            # если явно 'mismatch' — можно пропустить; если хочешь отображать как «неуязвим», можно включить
            if vcheck == 'mismatch':
                continue

            entry = matches.get(vid)
            if entry is None:
                entry = {'vuln': v, 'software': []}
                matches[vid] = entry

            entry['software'].append(
                {
                    'row': sw,
                    'version_check': vcheck,
                }
            )

    return matches


# ---------------------------------------------------------------------------
# Генерация Word-отчёта
# ---------------------------------------------------------------------------

def generate_vulnerability_report(output_dir: Path | None = None) -> Path:
    """
    Генерирует отчёт в формате .docx.
    Возвращает путь к созданному файлу.
    """
    agents, soft_by_agent, vulns = _load_data()

    if output_dir is None:
        output_dir = Path(RESULTS_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    report_path = output_dir / f'vuln_report_{timestamp}.docx'

    doc = Document()
    doc.add_heading('Отчёт по уязвимостям', level=1)
    doc.add_paragraph(datetime.now().strftime('Дата формирования: %d.%m.%Y %H:%M:%S'))

    if not agents:
        doc.add_paragraph('Нет данных о проверяемых хостах.')
        doc.save(report_path)
        return report_path

    total_matches = 0
    host_labels: List[str] = []

    for agent in agents:
        hostname = agent['hostname'] or agent['agent_id']
        ip = agent['ip_address'] or ''
        host_label = f'{hostname} ({ip})' if ip else hostname
        host_labels.append(host_label)

        matches = match_vulns_for_agent(agent, soft_by_agent, vulns)
        if not matches:
            continue

        total_matches += len(matches)
        doc.add_heading(f'Хост: {host_label}', level=2)
        doc.add_paragraph('Обнаружены потенциальные уязвимости:')

        for entry in matches.values():
            v = entry['vuln']
            sw_list = entry['software']

            cve = v['cve'] or ''
            title = v['title'] or ''
            vendor = v['vendor'] or ''
            product = v['product'] or ''
            severity = v['severity'] or ''
            cvss = v['cvss'] or ''
            url = v['url'] or ''

            p = doc.add_paragraph(style='List Bullet')

            base_text = f'{cve} – {title}'
            if product or vendor:
                base_text += f' (продукт: {product or "не указан"}, вендор: {vendor or "не указан"})'
            if severity or cvss:
                base_text += f' [критичность: {severity or "не указана"}, CVSS: {cvss or "n/a"}]'
            if url:
                base_text += f' | Подробнее: {url}'

            run = p.add_run(base_text)

            # перечисляем ПО на хосте, для которого мы решили, что оно может быть уязвимо
            for item in sw_list:
                sw = item['row']
                vcheck = item['version_check']
                sw_name = sw['name'] or ''
                sw_ver = sw['version'] or ''

                detail = f'\n  - ПО на хосте: {sw_name}'
                if sw_ver:
                    detail += f' (установленная версия: {sw_ver})'

                if vcheck == 'match':
                    detail += ' — установленная версия входит в диапазон уязвимых.'
                elif vcheck == 'unknown':
                    detail += ' — версия уязвимости не определена однозначно (требуется ручная проверка).'

                p.add_run(detail)

    doc.add_heading('Итог', level=2)
    if total_matches == 0:
        doc.add_paragraph('Уязвимости в системе не были обнаружены.')
    else:
        doc.add_paragraph(f'Всего найдено потенциально релевантных уязвимостей: {total_matches}.')

    uniq_hosts = ', '.join(sorted(set(host_labels)))
    doc.add_paragraph(f'Проверяемые хосты: {uniq_hosts}')

    doc.save(report_path)
    return report_path
